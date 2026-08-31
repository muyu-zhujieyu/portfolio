# -*- coding: utf-8 -*-
"""
液压伺服阀样本分析 - 20201010 文档解析服务

功能:
  1. 自动查找 D:/kg0623 下包含 20201010 的 docx 文件
  2. 提取文档段落文本、原始图片
  3. 将图片与前后文本上下文匹配
  4. 生成样本清单 sample_manifest_20201010.json
"""
import os
import json
import zipfile
import glob
import io
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple

BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
SAMPLE_DIR = os.path.join(BASE_DIR, "data", "sample_analysis")
RAW_IMAGES_DIR = os.path.join(SAMPLE_DIR, "raw_images")
ANALYSIS_OUTPUTS_DIR = os.path.join(SAMPLE_DIR, "analysis_outputs")
OVERLAYS_DIR = os.path.join(SAMPLE_DIR, "overlays")

# 部位关键词匹配规则
PART_KEYWORDS = [
    ("气隙垫片", ["气隙垫片", "气隙", "垫片", "间隙"]),
    ("马达螺钉", ["马达螺钉", "马达", "螺钉", "N.M"]),
    ("上壳体螺钉", ["上壳体螺钉", "上壳体"]),
    ("上壳体回油螺钉", ["上壳体回油螺钉", "回油螺钉", "回油"]),
    ("衔铁组件", ["衔铁", "衔铁组件"]),
    ("喷嘴挡板", ["喷嘴", "挡板", "喷嘴挡板"]),
    ("阀芯阀套", ["阀芯", "阀套", "滑阀"]),
    ("力矩马达", ["力矩马达", "马达"]),
    ("线圈与磁路", ["线圈", "磁路"]),
    ("密封组件", ["密封"]),
    ("反馈杆", ["反馈杆", "反馈"]),
]


class SampleDocReaderService:

    def __init__(self):
        os.makedirs(RAW_IMAGES_DIR, exist_ok=True)
        os.makedirs(ANALYSIS_OUTPUTS_DIR, exist_ok=True)
        os.makedirs(OVERLAYS_DIR, exist_ok=True)
        self._doc_path = ""
        self._find_doc()

    def _find_doc(self):
        """自动查找包含 20201010 的 docx 文件"""
        patterns = [
            os.path.join(BASE_DIR, "【公开】20201010.docx"),
            os.path.join(BASE_DIR, "[公开]20201010.docx"),
            os.path.join(BASE_DIR, "20201010.docx"),
        ]
        for p in patterns:
            if os.path.exists(p):
                self._doc_path = p
                return
        for f in glob.glob(os.path.join(BASE_DIR, "*20201010*.docx")):
            self._doc_path = f
            return

    # ================================================================
    # 公开接口
    # ================================================================

    def extract_doc(self) -> Dict[str, Any]:
        """解析 20201010 docx，提取原始图片和上下文，生成样本清单"""
        if not self._doc_path or not os.path.exists(self._doc_path):
            return {"状态": "失败", "错误": f"未找到 20201010 docx 文件"}

        # 1. 提取文本段落（按顺序，含位置）
        paragraphs = self._extract_paragraphs()

        # 2. 提取所有图片并保留顺序
        images = self._extract_images()

        # 3. 将图片与前后文本上下文关联
        samples = self._match_images_to_context(paragraphs, images)

        # 4. 将图片归类到部位
        samples = self._classify_samples(samples)

        # 5. 保存样本清单
        manifest_path = os.path.join(SAMPLE_DIR, "sample_manifest_20201010.json")
        manifest = {
            "meta": {
                "生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "源文件": os.path.basename(self._doc_path),
                "源文件路径": self._doc_path,
                "图片总数": len(images),
                "样本总数": len(samples),
                "原始图片目录": RAW_IMAGES_DIR,
                "说明": "所有原始图片来自 20201010 文档，未经修改",
            },
            "样本列表": samples,
        }
        with open(manifest_path, "w", encoding="utf-8") as f:
            json.dump(manifest, f, ensure_ascii=False, indent=2)

        return {
            "状态": "成功",
            "源文件": os.path.basename(self._doc_path),
            "提取图片数": len(images),
            "生成样本数": len(samples),
            "原始图片保存路径": RAW_IMAGES_DIR,
            "样本清单保存路径": manifest_path,
            "说明": "所有原始曲线图片来自 20201010 文档提取，未经修改",
        }

    def get_file_info(self) -> Dict[str, Any]:
        """返回文件信息"""
        exists = os.path.exists(self._doc_path) if self._doc_path else False
        manifest_path = os.path.join(SAMPLE_DIR, "sample_manifest_20201010.json")
        manifest_exists = os.path.exists(manifest_path)

        # 统计 raw_images 中的图片
        raw_count = 0
        if os.path.exists(RAW_IMAGES_DIR):
            raw_count = len([f for f in os.listdir(RAW_IMAGES_DIR)
                             if f.endswith(('.png', '.jpg', '.jpeg'))])

        # 统计样本数
        sample_count = 0
        if manifest_exists:
            with open(manifest_path, encoding="utf-8") as f:
                data = json.load(f)
                sample_count = data.get("meta", {}).get("样本总数", len(data.get("样本列表", [])))

        return {
            "文件名": os.path.basename(self._doc_path) if self._doc_path else "—",
            "文件路径": self._doc_path,
            "文件大小": os.path.getsize(self._doc_path) if exists else 0,
            "文件大小_可读": self._fmt_size(os.path.getsize(self._doc_path)) if exists else "—",
            "解析状态": "已解析" if manifest_exists else "待解析",
            "提取图片数": raw_count,
            "样本总数": sample_count,
            "原始图片目录": RAW_IMAGES_DIR,
            "说明": "曲线来源：20201010 原始文档提取，未修改原始曲线。",
        }

    def get_manifest(self) -> Dict[str, Any]:
        """返回完整样本清单"""
        manifest_path = os.path.join(SAMPLE_DIR, "sample_manifest_20201010.json")
        if not os.path.exists(manifest_path):
            return {"状态": "失败", "错误": "样本清单不存在，请先执行 POST /api/sample-analysis/extract-doc"}

        with open(manifest_path, encoding="utf-8") as f:
            return json.load(f)

    # ================================================================
    # 内部方法
    # ================================================================

    def _extract_paragraphs(self) -> List[Dict[str, Any]]:
        """提取文档所有段落文本"""
        try:
            from docx import Document
            doc = Document(self._doc_path)
            paragraphs = []
            for i, p in enumerate(doc.paragraphs):
                t = p.text.strip()
                if t:
                    paragraphs.append({
                        "段落索引": i,
                        "段落文本": t,
                        "字符数": len(t),
                    })
            return paragraphs
        except Exception:
            return []

    def _extract_images(self) -> List[Dict[str, Any]]:
        """从 docx 中提取所有图片并按出现顺序排列"""
        images = []

        # 利用 python-docx 的 inline_shapes 获取图片在段落中的位置
        try:
            from docx import Document
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            doc = Document(self._doc_path)

            # 遍历段落，找到包含图片的段落
            for pi, para in enumerate(doc.paragraphs):
                # 遍历段落的 runs，检查是否有 drawing 元素（图片）
                for ri, run in enumerate(para.runs):
                    # 检查是否有 inline shape
                    if hasattr(run, '_r') and run._r is not None:
                        drawings = run._r.findall(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
                        )
                        if not drawings:
                            drawings = run._r.findall(
                                '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
                            )
                        for _ in drawings:
                            # 通过关系找到对应的图片
                            for rel_id, rel in doc.part.rels.items():
                                target = rel.target_ref
                                if 'image' in str(rel.reltype).lower() and target not in [
                                    img.get('target') for img in images if img.get('target')
                                ]:
                                    images.append({
                                        'target': target,
                                        'rel_id': rel_id,
                                        '段落索引': pi,
                                        '段落文本': para.text.strip()[:120],
                                        '图片序号': len(images) + 1,
                                    })
                                    break

            # 如果上述方法未找到图片，直接用 zip 方式提取所有图片
            if not images:
                with zipfile.ZipFile(self._doc_path, 'r') as z:
                    img_files = sorted([
                        f for f in z.namelist()
                        if 'image' in f.lower() and f.endswith('.png')
                    ])
                    for idx, imgf in enumerate(img_files):
                        # 获取图片在 zip 中的顺序作为其文档位置
                        images.append({
                            'target': imgf,
                            '图片序号': idx + 1,
                            '段落索引': -1,
                            '段落文本': '',
                        })

            # 保存所有图片到 raw_images
            with zipfile.ZipFile(self._doc_path, 'r') as z:
                for idx, img_info in enumerate(images):
                    target = img_info.get('target', '')
                    if not target or target not in z.namelist():
                        # 尝试通过 rel 查找图片
                        found = False
                        for name in z.namelist():
                            if 'image' in name.lower() and target and name.endswith(
                                    os.path.splitext(target)[-1] or '.png'):
                                target = name
                                found = True
                                break
                        if not found:
                            continue

                    ext = os.path.splitext(target)[1] or '.png'
                    out_name = f"sample_{idx + 1:03d}_original{ext}"
                    out_path = os.path.join(RAW_IMAGES_DIR, out_name)

                    if not os.path.exists(out_path):
                        with z.open(target) as src:
                            with open(out_path, 'wb') as dst:
                                dst.write(src.read())

                    img_info['原始图片路径'] = out_path
                    img_info['原始图片文件名'] = out_name

            return images

        except Exception as e:
            print(f"Image extraction error: {e}")
            return []

    def _match_images_to_context(self, paragraphs: List[Dict],
                                   images: List[Dict]) -> List[Dict]:
        """将图片与前后文本上下文关联"""
        samples = []

        for img in images:
            pi = img.get('段落索引', -1)

            # 查找前后上下文段落
            prev_context = []
            next_context = []

            if pi >= 0 and paragraphs:
                # 前 3 个段落
                for p in paragraphs:
                    if 0 <= p['段落索引'] < pi:
                        prev_context.append(p)
                prev_context = prev_context[-3:]

                # 后 3 个段落
                for p in paragraphs:
                    if p['段落索引'] > pi:
                        next_context.append(p)
                        if len(next_context) >= 3:
                            break

            # 生成上下文摘要
            prev_text = "；".join([p['段落文本'][:80] for p in prev_context])
            next_text = "；".join([p['段落文本'][:80] for p in next_context])
            context_summary = f"前: {prev_text[:100]} | 后: {next_text[:100]}"

            sample = {
                "样本编号": f"S{img.get('图片序号', 0):03d}",
                "样本名称": f"样本 {img.get('图片序号', 0):03d}",
                "原始图片路径": img.get('原始图片路径', ''),
                "原始图片文件名": img.get('原始图片文件名', ''),
                "原始文档文件名": os.path.basename(self._doc_path),
                "前置上下文": [p['段落文本'] for p in prev_context],
                "后置上下文": [p['段落文本'] for p in next_context],
                "上下文摘要": context_summary,
                "部位名称": "其他",  # 后续分类
                "是否已分析": False,
            }
            samples.append(sample)

        return samples

    def _classify_samples(self, samples: List[Dict]) -> List[Dict]:
        """根据上下文将样本归类到部位"""
        for sample in samples:
            all_text = " ".join(
                sample.get('前置上下文', []) + sample.get('后置上下文', [])
            )
            sample['部位名称'] = self._classify_part(all_text)
        return samples

    def _classify_part(self, text: str) -> str:
        """根据文本内容归类部位"""
        scores = {}
        for part_name, keywords in PART_KEYWORDS:
            score = sum(1 for kw in keywords if kw in text)
            if score > 0:
                scores[part_name] = score

        if not scores:
            return "其他"

        # 返回得分最高的部位
        return max(scores, key=scores.get)

    @staticmethod
    def _fmt_size(size_bytes: int) -> str:
        if size_bytes < 1024: return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024: return f"{size_bytes / 1024:.1f} KB"
        else: return f"{size_bytes / 1024 / 1024:.1f} MB"


# 单例
doc_reader = SampleDocReaderService()
